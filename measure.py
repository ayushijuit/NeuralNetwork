import docx
from nltk.tokenize import sent_tokenize

doc=docx.Document('C:/Users/acer/Desktop/final yr project/Doc.docx')
wholedoc=""
for para in doc.paragraphs:
    wholedoc=wholedoc+para.text
summary_sentences = sent_tokenize(wholedoc)

testdocclass1=docx.Document('C:/Users/acer/Desktop/final yr project/test.docx')
wholedoc=""
for para in testdocclass1.paragraphs:
    wholedoc=wholedoc+para.text
test_sentences = sent_tokenize(wholedoc)

tp=0
fp=0
fn=0
tn=0
for x in range(len(summary_sentences)):
    if (summary_sentences[x] in test_sentences):
        tp=tp+1
    else:
        fp=fp+1
pre=tp/(tp+fp)*100
print("Precision:",pre)

for x in range(len(test_sentences)):
    if(test_sentences[x] not in summary_sentences):
        fn=fn+1
rec=tp/(tp+fn)*100
print("Recall:",rec)

fscore=2*((pre*rec)/(pre+rec))
print("F-score:",fscore)
print(tp)
print(fn)

